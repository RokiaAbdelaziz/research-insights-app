import io
import re
import json
import uuid
import base64
from pathlib import Path
from datetime import date, datetime

import streamlit as st
import anthropic
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

st.set_page_config(
    page_title="Research & Insights Automation Engine",
    page_icon="assets/logo_placeholder.png",
    layout="wide"
)

st.title("🔍 Research & Insights Specialist Engine")
st.caption("Automated Research & Business Intelligence Generator")
# ----------------------------
# Branding
# ----------------------------
# Swap assets/logo_placeholder.png and assets/favicon.png for real brand
# files any time — same filenames, no code changes needed. The circular "3M"
# mark is a placeholder standing in for an actual logo.
try:
    st.logo("assets/logo_placeholder.png", size="large")
except Exception:
    pass  # older Streamlit versions without st.logo() just skip this gracefully
 
st.markdown("""
<style>
/* Tighter, more deliberate type for a premium dark-theme feel */
h1, h2, h3 { letter-spacing: 0.02em; }
 
/* Accent underline on the active tab instead of the default color block */
.stTabs [data-baseweb="tab-list"] { gap: 4px; }
.stTabs [aria-selected="true"] {
    border-bottom: 2px solid #E5E5E5 !important;
    color: #FFFFFF !important;
}
 
/* Slightly muted card-like look for expanders (Knowledge Base entries) */
[data-testid="stExpander"] {
    border: 1px solid #2A2A2A;
    border-radius: 8px;
}
 
/* Buttons: subtle border instead of flat fill, feels more editorial */
.stButton > button, .stDownloadButton > button, .stLinkButton > a {
    border: 1px solid #E5E5E5 !important;
    background-color: transparent !important;
    color: #F5F5F5 !important;
}
.stButton > button:hover, .stDownloadButton > button:hover, .stLinkButton > a:hover {
    background-color: #E5E5E5 !important;
    color: #000000 !important;
}
</style>
""", unsafe_allow_html=True)
 
col_logo, col_title = st.columns([1, 8])
with col_logo:
    st.image("assets/logo_placeholder.png", width=64)
with col_title:
    st.title("Research & Insights Specialist Engine")
    st.caption("Automated Research & Business Intelligence Generator — aligned to the Research & Insights Specialist Handbook")
 
# ----------------------------
# Configuration
# ----------------------------
try:
    api_key = st.secrets["ANTHROPIC_API_KEY"]
except Exception:
    st.error("⚠️ API Key not found in Streamlit Secrets. Please check your app settings.")
    st.stop()

MODEL_NAME = "claude-sonnet-5"

KB_CATEGORIES = ["Market", "Consumer", "Competitors", "Content", "Trends", "Technology", "Statistics"]

FOCUS_TO_CATEGORY = {
    "Consumer Behavior & Pain Points": "Consumer",
    "Competitor Positioning & Branding": "Competitors",
    "Packaging & Visual Identity Trends": "Trends",
    "Social Media & Platform Updates": "Content",
    "AI Tools & MarTech Innovations": "Technology",
    "Market Gaps & Business Opportunities": "Market",
}

DELIVERABLE_TYPES = [
    "General Research Summary",
    "Competitor Research Report",
    "Industry Update Report",
    "Marketing & AI Update",
    "Competitor Highlights",
    "Opportunity Report",
    "Trend Collection & Opportunity Audit",
    "Client Onboarding Research File",
    "Knowledge Base Review",
]

# ----------------------------
# Knowledge Base — Google Apps Script bridge backend
# ----------------------------
# PRIMARY ARCHIVE: every generated report is created as its own real,
# editable Google Doc inside a Drive folder you choose. Instead of a
# service account + Docs/Drive API (which commonly fails on personal Google
# accounts due to service accounts having no Drive storage quota, and
# folder-sharing permissions not propagating the way people expect), this
# talks to a small Google Apps Script Web App that runs under YOUR OWN
# Google account. Docs it creates use your normal Drive quota and are
# automatically yours.
#
# SETUP (see apps_script_bridge.gs for the full script + inline instructions):
#   1. Paste apps_script_bridge.gs into a new project at script.google.com.
#   2. Set your own SHARED_SECRET string inside it.
#   3. Deploy as a Web App (Execute as: Me, Who has access: Anyone) and
#      authorize it.
#   4. Put these into Streamlit secrets:
#        GOOGLE_APPS_SCRIPT_URL    = "<the web app URL>"
#        GOOGLE_APPS_SCRIPT_SECRET = "<the same secret you set in the script>"
#        GOOGLE_KB_FOLDER_ID       = "<Drive folder ID to use as the KB>"
#
# FALLBACK: if these secrets aren't set, or a call to the script fails, the
# app still saves a local copy (knowledge_base/ on disk) so nothing is lost
# — but that local copy won't persist on ephemeral hosting and isn't the
# shared team archive.
APPS_SCRIPT_URL = st.secrets.get("GOOGLE_APPS_SCRIPT_URL")
APPS_SCRIPT_SECRET = st.secrets.get("GOOGLE_APPS_SCRIPT_SECRET")
KB_FOLDER_ID = st.secrets.get("GOOGLE_KB_FOLDER_ID")


def google_kb_configured():
    return bool(APPS_SCRIPT_URL and APPS_SCRIPT_SECRET and KB_FOLDER_ID)


def _call_apps_script(action, payload, timeout=60):
    """Returns (result_dict, error_string). Exactly one will be None."""
    if not google_kb_configured():
        return None, "Google Knowledge Base not configured in secrets."
    data = {"action": action, "secret": APPS_SCRIPT_SECRET}
    data.update(payload)
    try:
        resp = requests.post(APPS_SCRIPT_URL, json=data, timeout=timeout)
        resp.raise_for_status()
        result = resp.json()
        if isinstance(result, dict) and result.get("error"):
            return None, result["error"]
        return result, None
    except requests.exceptions.RequestException as e:
        return None, f"Request to Apps Script failed: {e}"
    except ValueError:
        return None, "Apps Script returned a non-JSON response (check the deployment is a Web App, not an editor URL)."


def _split_bold_segments(text):
    """Returns [(substring, is_bold), ...] for a line containing **bold** markers."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    segments = []
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith('**') and part.endswith('**')
        segments.append((part[2:-2] if is_bold else part, is_bold))
    return segments


def markdown_to_blocks(markdown_text: str, rtl: bool = False):
    """
    Converts the app's constrained Markdown subset into a list of simple
    JSON-serializable blocks the Apps Script bridge knows how to render as
    native Google Docs formatting (headings, bold runs, bullet/numbered
    lists, plain paragraphs). Markdown tables become bold-header
    pipe-separated rows rather than native Docs tables, to keep this
    single-pass and dependency-free on the Apps Script side.
    """
    blocks = []
    lines = markdown_text.splitlines()
    i, n = 0, len(lines)

    while i < n:
        raw_line = lines[i].rstrip()

        if not raw_line.strip():
            i += 1
            continue

        if re.match(r'^-{3,}$', raw_line.strip()):
            i += 1
            continue

        header_match = re.match(r'^(#{1,4})\s+(.*)', raw_line)
        if header_match:
            level = len(header_match.group(1))
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', header_match.group(2).strip())
            blocks.append({"type": "heading", "level": level, "text": text, "rtl": rtl})
            i += 1
            continue

        if raw_line.strip().startswith('|'):
            table_lines = []
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [
                [c.strip() for c in tl.strip('|').split('|')]
                for tl in table_lines
                if not re.match(r'^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?$', tl)
            ]
            for r_idx, row_cells in enumerate(rows):
                blocks.append({"type": "table_row", "cells": row_cells, "header": r_idx == 0, "rtl": rtl})
            continue

        bullet_match = re.match(r'^[\-\*]\s+(.*)', raw_line)
        numbered_match = re.match(r'^\d+\.\s+(.*)', raw_line)
        if bullet_match or numbered_match:
            content = (bullet_match or numbered_match).group(1)
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)
            blocks.append({
                "type": "numbered" if numbered_match else "bullet",
                "text": content,
                "rtl": rtl
            })
            i += 1
            continue

        runs = [{"text": t, "bold": b} for t, b in _split_bold_segments(raw_line.strip())]
        blocks.append({"type": "paragraph", "runs": runs, "rtl": rtl})
        i += 1

    return blocks


def create_kb_google_doc(title, folder_id, report_text, rtl, properties):
    blocks = markdown_to_blocks(report_text, rtl=rtl)
    result, error = _call_apps_script("create_doc", {
        "title": title[:200],
        "folderId": folder_id,
        "blocks": blocks,
        "properties": properties
    })
    if error:
        raise RuntimeError(error)
    return result["docId"], result["url"]


def list_kb_google_docs(folder_id):
    result, error = _call_apps_script("list_docs", {"folderId": folder_id})
    if error:
        raise RuntimeError(error)
    return result.get("files", [])


def export_google_doc_as_docx(doc_id):
    result, error = _call_apps_script("export_docx", {"docId": doc_id})
    if error:
        raise RuntimeError(error)
    return base64.b64decode(result["base64"])


# ----------------------------
# Local Knowledge Base fallback (file-based)
# ----------------------------
# Used only when the Google bridge isn't configured or a call to it fails,
# so a report is never silently lost. Does NOT persist on ephemeral hosting.
KB_DIR = Path("knowledge_base")
KB_FILES_DIR = KB_DIR / "files"
KB_INDEX_PATH = KB_DIR / "index.json"


def _ensure_kb_dirs():
    KB_FILES_DIR.mkdir(parents=True, exist_ok=True)


def load_local_kb_index():
    _ensure_kb_dirs()
    if not KB_INDEX_PATH.exists():
        return []
    try:
        return json.loads(KB_INDEX_PATH.read_text(encoding="utf-8"))
    except Exception:
        return []


def save_local_kb_index(index):
    _ensure_kb_dirs()
    KB_INDEX_PATH.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")


def archive_to_local_knowledge_base(entry_meta, report_text, docx_bytes):
    _ensure_kb_dirs()
    entry_id = entry_meta["id"]

    md_path = KB_FILES_DIR / f"{entry_id}.md"
    md_path.write_text(report_text, encoding="utf-8")

    docx_path = None
    if docx_bytes is not None:
        docx_path = KB_FILES_DIR / f"{entry_id}.docx"
        docx_path.write_bytes(docx_bytes)

    entry_meta["md_path"] = str(md_path)
    entry_meta["docx_path"] = str(docx_path) if docx_path else None

    index = load_local_kb_index()
    index.append(entry_meta)
    save_local_kb_index(index)
    return entry_meta


# ----------------------------
# Report depth / length controls
# ----------------------------
with st.sidebar:
    st.header("⚙️ Report Settings")
    max_output_tokens = st.slider(
        "Max output tokens", min_value=4000, max_value=16000, value=12000, step=1000,
        help="A full 9-section handbook report typically needs 6,000-10,000 tokens. "
             "Raise this if reports keep getting cut off."
    )
    max_searches = st.slider(
        "Max web searches allowed", min_value=5, max_value=30, value=20, step=5,
        help="Comprehensive market/competitor reports need 10-20+ searches. "
             "Too low a cap forces shallow, under-researched output."
    )
    st.markdown("---")
    st.subheader("📚 Knowledge Base status")
    if google_kb_configured():
        st.success("Google Docs Knowledge Base is configured.")
    else:
        st.warning(
            "Google Docs Knowledge Base not configured — reports will only "
            "be saved locally (not persistent on ephemeral hosting). "
            "See apps_script_bridge.gs for setup."
        )

# ----------------------------
# Tabs: Generate vs. Knowledge Base
# ----------------------------
tab_generate, tab_kb = st.tabs(["📝 Generate Report", "📚 Knowledge Base"])

# ============================================================
# TAB 1 — GENERATE REPORT
# ============================================================
with tab_generate:
    with st.form("comprehensive_research_form"):
        col1, col2 = st.columns(2)

        with col1:
            project_name = st.text_input(
                "1. Project / Client Name",
                placeholder="e.g., Specialty Coffee Retail Audit 2026"
            )

            industry = st.selectbox(
                "2. Core Industry Focus",
                ["Fashion", "Interior & Architecture", "Food & Beverage (F&B)", "Medical", "Cross-Industry / General"]
            )

            deliverable_type = st.selectbox(
                "3. Deliverable Type",
                DELIVERABLE_TYPES
            )

        with col2:
            language = st.selectbox(
                "4. Output Language & Tone",
                [
                    "English (Formal Business / Research)",
                    "Modern Standard Arabic (فصحى)",
                    "Egyptian Natural Language Arabic (عامية مصرية احترافية)"
                ]
            )

            priority_focus = st.multiselect(
                "5. Primary Focus Areas (Select up to 3)",
                list(FOCUS_TO_CATEGORY.keys()),
                default=["Consumer Behavior & Pain Points", "Market Gaps & Business Opportunities"]
            )

            target_market = st.text_input(
                "6. Target Market / Geographic Region",
                placeholder="e.g., Egypt, GCC Region, Global"
            )

        objective = st.text_area(
            "7. Research Objective & Business Question (Crucial)",
            placeholder="Why are we conducting this research? What specific business question or decision will this output support? Who will use it?",
            height=100
        )

        specific_questions = st.text_area(
            "8. Key Research Questions (Optional)",
            placeholder="List specific questions to answer (e.g., What are top competitors charging? What packaging material is trending?)",
            height=80
        )

        kb_categories = st.multiselect(
            "9. Knowledge Base Categories (for archiving)",
            KB_CATEGORIES,
            default=sorted({FOCUS_TO_CATEGORY[f] for f in priority_focus if f in FOCUS_TO_CATEGORY}) or ["Market"]
        )

        submit_button = st.form_submit_button("Generate Research Report")

    # ----------------------------
    # System Prompt — mirrors the Handbook's exact structure & standards
    # ----------------------------
    COMPREHENSIVE_SYSTEM_PROMPT = f"""
You are an automated Research & Insights Specialist operating under this company's Research & Insights Specialist Handbook. Your mission is to reduce uncertainty before decisions are made — not to make the decisions yourself. You provide reliable research, verified facts, meaningful insights, and valuable opportunities so the team can move with confidence instead of assumptions.

Today's date is {date.today().strftime('%B %d, %Y')}. Use this to judge what counts as current, and to phrase search queries correctly.

CORE MISSION (from the Handbook):
- Find reliable information. Verify its accuracy. Understand what it means. Identify opportunities. Organize knowledge.
- You are NOT expected to make strategic decisions. You ARE expected to make those decisions easier by explaining WHAT is happening, WHY it is happening, WHY it matters, and WHERE the opportunity exists.
- Do NOT design creative assets, write final social media copy, build full brand strategies, or plan content campaigns — that is out of scope. Leave execution to the creative/strategy teams. If your draft starts reading like a content plan instead of a research report, rewrite it as analysis, not as production instructions.

WORKFLOW YOU MUST FOLLOW INTERNALLY BEFORE WRITING (Handbook Steps 1-9):
1. Understand the objective given below — why this research, what decision it supports, who will use it.
2. Break it into specific research questions (don't research the topic broadly).
3. Collect information from MULTIPLE independent sources — never rely on a single site for any claim that matters.
4. Verify each piece of information: Is it accurate? Is it still relevant/current? Confirmed by another source? Does it apply to the target market specified? Explicitly check the publication date of every source you cite, and flag anything more than ~18 months old as potentially dated. Where possible, trace a claim back to its original/primary source rather than resting on a secondary summary of it.
5. Organize findings into the Handbook's topical categories: Market, Consumer, Competitors, Content, Trends, Technology, Statistics. A single finding can belong to more than one.
6. Extract insights — after every important finding, ask "what does this mean for our business?"
7. Identify opportunities — what gap, pain point, or advantage does this create?
8. Build recommendations that are a natural conclusion of the findings above, never an assumption stated before the evidence supporting it.
9. (Handled outside your output — the application archives your finished report automatically.)

RESEARCH DEPTH & RIGOR:
- This is a paid, professional deliverable. Run as many distinct web searches as needed (rarely fewer than 8-10 for a full report) — do not stop after 1-2. Search each sub-question separately rather than combining topics into one broad query.
- Prefer concrete numbers, percentages, dates, and named entities over vague generalizations ("growing rapidly," "consumers increasingly prefer"). Every FACT must be traceable to a specific, named source.
- Where sources disagree, state the discrepancy rather than silently picking one.
- Never quote more than a short phrase (under ~15 words) from any single source, and never more than one such phrase per source — paraphrase everything else in your own words.

SOURCE PRIORITY (Handbook standard — always prefer the highest tier available):
- Priority 1: Official platforms, official documentation, peer-reviewed research papers, government publications.
- Priority 2: Industry reports and trusted research organizations.
- Priority 3: Trusted marketing and business publications.
- Priority 4: Industry experts.
- Priority 5: Communities and discussions (consumer-sentiment signal only — never the sole basis for a Fact).

CLASSIFICATION FRAMEWORK (STRICT SEPARATION — every substantive statement must be tagged as exactly one of these):
- FACT: A verified piece of information supported by reliable, cited sources.
- OBSERVATION: A noticeable pattern or repeated behavior identified through the research.
- INSIGHT: An explanation of WHY something matters and how it affects the business — not just what happened.
- OPPORTUNITY: A possible area where the company or client can create value.
- RECOMMENDATION: A suggested, evidence-backed action — never personal opinion, never stated before the findings that justify it.

PRIORITIZATION (do not present every finding as equally important):
- For every Insight, Opportunity, and Recommendation, assign a priority of [High], [Medium], or [Low] based on business impact and urgency, and lead each subsection with the [High] items first.
- In the Executive Summary, explicitly call out the single highest-priority finding and the single highest-priority recommendation — the reader should know in 10 seconds what matters most.

REQUIRED DELIVERABLE STRUCTURE — this is the Handbook's Research Standard. Use exactly these section headers, in exactly this order, nothing added or removed:
1. Research Objective
2. Research Questions
3. Executive Summary
4. Key Findings — organized under topical subheadings from this exact set: Market, Consumer, Competitors, Content, Trends, Technology, Statistics (omit any category with no findings). Within each subheading, tag every statement inline with its classification in bold brackets, e.g. "**[Fact — Priority 2]** ..." or "**[Insight — High]** ...".
5. Supporting Data — a Markdown table of the key quantitative data points referenced above (columns: Data Point, Value, Source, Source Tier, Publication Date).
6. Key Insights — the strategic "why it matters" analysis, each tagged with a priority.
7. Opportunities — each tagged with a priority.
8. Recommendations — each tagged with a priority, and each one traceable to a specific finding or insight above it.
9. References — every source actually used, grouped under "Priority 1" through "Priority 5" headers, with source name, URL, and the publication date you found for it (write "date not stated" if unavailable).

FORMATTING:
- Output ONLY the finished report in clean Markdown — no preamble like "Here is your report."
- End with this exact note: "This report covers WHAT is happening, WHY it matters, and WHERE the opportunity is. Creative execution (design, copywriting, campaign planning) is out of scope and left to the strategy/creative team."

LANGUAGE & TONE:
- IF Language = English (Formal Business / Research): formal, sharp, authoritative agency-research English.
- IF Language = Modern Standard Arabic (فصحى): precise, standard corporate Arabic, clear analytical terminology.
- IF Language = Egyptian Natural Language Arabic (عامية مصرية احترافية): professional, clean Egyptian Arabic suitable for local agency teams (عامية مصرية راقية ومفهومة لفرق العمل). Keep technical marketing/research terms intact in English (Insights, Conversion, Benchmarks, Target Audience, Positioning, CAGR, SKU, Fact, Observation, Opportunity, Recommendation) while making the surrounding prose read naturally, not like a translation. Section headers and body text should all be in Egyptian Arabic.
"""

    # ----------------------------
    # Helper: run one generation, continuing automatically if truncated
    # ----------------------------
    def generate_report(client, user_prompt, system_prompt, max_tokens, max_searches_n, status):
        messages = [{"role": "user", "content": user_prompt}]
        full_text = ""
        searches_seen = []

        for turn in range(4):  # hard safety cap on continuation turns
            status.update(label=f"Researching and drafting (pass {turn + 1})...")
            response = client.messages.create(
                model=MODEL_NAME,
                max_tokens=max_tokens,
                system=system_prompt,
                messages=messages,
                tools=[{
                    "type": "web_search_20250305",
                    "name": "web_search",
                    "max_uses": max_searches_n
                }],
            )

            for block in response.content:
                btype = getattr(block, "type", None)
                if btype == "server_tool_use" and getattr(block, "name", "") == "web_search":
                    query = getattr(block, "input", {}).get("query")
                    if query:
                        searches_seen.append(query)

            turn_text = "".join(
                block.text for block in response.content if getattr(block, "type", None) == "text"
            )
            full_text += turn_text

            if response.stop_reason != "max_tokens":
                break

            messages.append({"role": "assistant", "content": response.content})
            messages.append({
                "role": "user",
                "content": "Continue the report exactly where you left off. Do not repeat "
                           "any content already written, do not restart headers already "
                           "completed, and do not add any preamble."
            })

        return full_text, searches_seen

    # ----------------------------
    # Markdown -> Word (.docx) conversion (used for local backup / download)
    # ----------------------------
    def _set_paragraph_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_run_with_bold(paragraph, text, rtl=False):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if not part:
                continue
            is_bold = part.startswith('**') and part.endswith('**')
            content = part[2:-2] if is_bold else part
            run = paragraph.add_run(content)
            run.bold = is_bold
            if rtl:
                rPr = run._r.get_or_add_rPr()
                rtl_el = OxmlElement('w:rtl')
                rtl_el.set(qn('w:val'), '1')
                rPr.append(rtl_el)

    def _style_table_header_cell(cell):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:fill'), 'D9D9D9')
        cell._tc.get_or_add_tcPr().append(shading)

    def markdown_to_docx(markdown_text: str, rtl: bool = False, base_font: str = None) -> bytes:
        doc = Document()

        normal_style = doc.styles['Normal']
        if base_font:
            normal_style.font.name = base_font
            rPr = normal_style.element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:cs'), base_font)
        normal_style.font.size = Pt(11)

        if rtl:
            sectPr = doc.sections[0]._sectPr
            bidi = OxmlElement('w:bidi')
            sectPr.append(bidi)

        lines = markdown_text.splitlines()
        i = 0
        n = len(lines)

        while i < n:
            line = lines[i].rstrip()

            if not line.strip():
                i += 1
                continue

            if re.match(r'^-{3,}$', line.strip()):
                i += 1
                continue

            header_match = re.match(r'^(#{1,4})\s+(.*)', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2).strip()
                p = doc.add_heading(level=min(level, 4))
                _add_run_with_bold(p, text, rtl=rtl)
                if rtl:
                    _set_paragraph_rtl(p)
                i += 1
                continue

            if line.strip().startswith('|'):
                table_lines = []
                while i < n and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                rows = [
                    [c.strip() for c in tl.strip('|').split('|')]
                    for tl in table_lines
                    if not re.match(r'^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?$', tl)
                ]
                if rows:
                    ncols = len(rows[0])
                    table = doc.add_table(rows=0, cols=ncols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.style = 'Light Grid Accent 1'
                    for r_idx, row_cells in enumerate(rows):
                        row = table.add_row()
                        for c_idx in range(ncols):
                            cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                            cell = row.cells[c_idx]
                            cell.paragraphs[0].text = ""
                            p = cell.paragraphs[0]
                            _add_run_with_bold(p, cell_text, rtl=rtl)
                            if rtl:
                                _set_paragraph_rtl(p)
                            if r_idx == 0:
                                for run in p.runs:
                                    run.bold = True
                                _style_table_header_cell(cell)
                continue

            bullet_match = re.match(r'^[\-\*]\s+(.*)', line)
            if bullet_match:
                p = doc.add_paragraph(style='List Bullet')
                _add_run_with_bold(p, bullet_match.group(1), rtl=rtl)
                if rtl:
                    _set_paragraph_rtl(p)
                i += 1
                continue

            numbered_match = re.match(r'^\d+\.\s+(.*)', line)
            if numbered_match:
                p = doc.add_paragraph(style='List Number')
                _add_run_with_bold(p, numbered_match.group(1), rtl=rtl)
                if rtl:
                    _set_paragraph_rtl(p)
                i += 1
                continue

            p = doc.add_paragraph()
            _add_run_with_bold(p, line.strip(), rtl=rtl)
            if rtl:
                _set_paragraph_rtl(p)
            i += 1

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ----------------------------
    # Execution
    # ----------------------------
    if submit_button:
        if not objective:
            st.warning("⚠️ Please provide a Research Objective before generating.")
        else:
            client = anthropic.Anthropic(api_key=api_key)

            user_prompt = f"""
Execute research report generation based on the following standardized input:

- Project / Client Name: {project_name if project_name else "N/A"}
- Core Industry: {industry}
- Deliverable Type: {deliverable_type}
- Target Market: {target_market if target_market else "General / Unspecified"}
- Priority Focus Areas: {', '.join(priority_focus) if priority_focus else "None specified"}
- Preferred Language: {language}

BUSINESS OBJECTIVE:
{objective}

SPECIFIC RESEARCH QUESTIONS TO ANSWER:
{specific_questions if specific_questions else "Derive the key research questions yourself from the business objective above before researching."}
"""

            with st.status("Generating report...", expanded=True) as status:
                try:
                    report_text, searches_used = generate_report(
                        client, user_prompt, COMPREHENSIVE_SYSTEM_PROMPT,
                        max_output_tokens, max_searches, status
                    )
                    status.update(label="Report generated.", state="complete")
                except Exception as e:
                    status.update(label="Generation failed.", state="error")
                    st.error(f"Execution Error: {str(e)}")
                    st.stop()

            if searches_used:
                with st.expander(f"🔎 {len(searches_used)} web searches performed"):
                    for q in searches_used:
                        st.write(f"- {q}")

            st.success("Report Generated Successfully!")
            st.markdown("---")
            st.markdown(report_text)

            is_arabic = "Arabic" in language

            docx_bytes = None
            try:
                docx_bytes = markdown_to_docx(
                    report_text,
                    rtl=is_arabic,
                    base_font="Arial" if is_arabic else "Calibri"
                )
            except Exception as e:
                st.warning(f"Word export failed, but .txt/.md are still available below. ({e})")

            entry_id = str(uuid.uuid4())[:8]
            created_at = datetime.now().isoformat(timespec="seconds")
            properties = {
                "id": entry_id,
                "created_at": created_at,
                "project_name": project_name or "Untitled",
                "industry": industry,
                "deliverable_type": deliverable_type,
                "target_market": target_market or "General / Unspecified",
                "language": language,
                "priority_focus": ", ".join(priority_focus),
                "kb_categories": ", ".join(kb_categories),
                "objective": (objective or "")[:300],
            }
            title = f"{project_name or 'Untitled'} — {deliverable_type} ({date.today().isoformat()})"

            # Archive. Google Docs is the primary archive;
            # local save is a fallback so nothing is ever silently lost.
            if google_kb_configured():
                try:
                    doc_id, doc_url = create_kb_google_doc(
                        title, KB_FOLDER_ID, report_text, is_arabic, properties
                    )
                    st.success(
                        f"📚 Archived as an editable Google Doc: [{title}]({doc_url})"
                    )
                except Exception as e:
                    st.error(
                        f"⚠️ Could not archive to Google Docs ({e}). "
                        f"Saving a local copy instead so the report isn't lost."
                    )
                    archive_to_local_knowledge_base(
                        {**properties, "searches_used": searches_used}, report_text, docx_bytes
                    )
            else:
                archive_to_local_knowledge_base(
                    {**properties, "searches_used": searches_used}, report_text, docx_bytes
                )
                st.info(
                    f"📚 Saved locally under: {', '.join(kb_categories)} "
                    f"(Google Docs Knowledge Base not configured — see sidebar)."
                )

            # Action Buttons
            st.markdown("---")
            safe_name = (project_name or "Research").strip().replace(" ", "_")

            col_dl1, col_dl2, col_dl3 = st.columns(3)
            with col_dl1:
                st.download_button(
                    label="📄 Download Report (.txt)",
                    data=report_text,
                    file_name=f"{safe_name}_Report.txt",
                    mime="text/plain"
                )
            with col_dl2:
                st.download_button(
                    label="📝 Download Report (.md)",
                    data=report_text,
                    file_name=f"{safe_name}_Report.md",
                    mime="text/markdown"
                )
            with col_dl3:
                if docx_bytes is not None:
                    st.download_button(
                        label="📘 Download Report (.docx)",
                        data=docx_bytes,
                        file_name=f"{safe_name}_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

# ============================================================
# TAB 2 — KNOWLEDGE BASE (browse / filter / re-download archived research)
# ============================================================
with tab_kb:
    st.subheader("📚 Company Knowledge Base")
    st.caption(
        "Every generated report is archived here automatically. "
        "Future research should build on these instead of starting from zero."
    )

    google_entries = []
    google_error = None
    if google_kb_configured():
        try:
            google_entries = list_kb_google_docs(KB_FOLDER_ID)
        except Exception as e:
            google_error = str(e)

    local_entries = load_local_kb_index()

    if google_error:
        st.error(f"Couldn't load the Google Docs Knowledge Base: {google_error}")

    if not google_entries and not local_entries:
        st.info("No archived research yet. Generate a report in the first tab to populate the Knowledge Base.")
    else:
        if google_entries:
            st.markdown("### 📄 Google Docs (primary archive)")

            all_industries = sorted({e["properties"].get("industry", "Unknown") for e in google_entries})
            col_f1, col_f2, col_f3 = st.columns(3)
            with col_f1:
                filter_industry = st.multiselect("Filter by Industry", all_industries, key="g_industry")
            with col_f2:
                filter_category = st.multiselect("Filter by KB Category", KB_CATEGORIES, key="g_category")
            with col_f3:
                search_term = st.text_input("Search project name / objective", placeholder="e.g., cold brew", key="g_search")

            filtered = google_entries
            if filter_industry:
                filtered = [e for e in filtered if e["properties"].get("industry") in filter_industry]
            if filter_category:
                filtered = [
                    e for e in filtered
                    if any(c in e["properties"].get("kb_categories", "") for c in filter_category)
                ]
            if search_term:
                term = search_term.lower()
                filtered = [
                    e for e in filtered
                    if term in e["properties"].get("project_name", "").lower()
                    or term in e["properties"].get("objective", "").lower()
                ]

            st.write(f"**{len(filtered)}** of {len(google_entries)} archived Google Docs match your filters.")
            st.markdown("---")

            for entry in filtered:
                props = entry.get("properties", {})
                label = f"{props.get('project_name', entry['name'])} — {props.get('industry', '—')} — {entry['createdTime'][:10]}"
                with st.expander(label):
                    st.write(f"**Deliverable Type:** {props.get('deliverable_type', '—')}")
                    st.write(f"**Target Market:** {props.get('target_market', '—')}")
                    st.write(f"**KB Categories:** {props.get('kb_categories', '—')}")
                    st.write(f"**Objective:** {props.get('objective', '—')}")
                    st.link_button("📄 Open in Google Docs", entry["url"])

                    if st.button("📘 Prepare .docx download", key=f"prep_docx_{entry['id']}"):
                        try:
                            docx_bytes = export_google_doc_as_docx(entry["id"])
                            st.download_button(
                                "📘 Download .docx",
                                data=docx_bytes,
                                file_name=f"{entry['name']}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_docx_{entry['id']}"
                            )
                        except Exception as e:
                            st.error(f"Export failed: {e}")

        if local_entries:
            st.markdown("### 💾 Local fallback copies")
            st.caption("Only present if Google Docs archiving wasn't configured or failed for a given report.")
            for entry in sorted(local_entries, key=lambda e: e["created_at"], reverse=True):
                label = f"{entry.get('project_name', 'Untitled')} — {entry.get('industry', '—')} — {entry['created_at'][:10]}"
                with st.expander(label):
                    st.write(f"**Deliverable Type:** {entry.get('deliverable_type', '—')}")
                    st.write(f"**Target Market:** {entry.get('target_market', '—')}")
                    st.write(f"**KB Categories:** {entry.get('kb_categories', '—')}")
                    st.write(f"**Objective:** {entry.get('objective', '—')}")

                    md_path = Path(entry["md_path"]) if entry.get("md_path") else None
                    docx_path = Path(entry["docx_path"]) if entry.get("docx_path") else None

                    bcol1, bcol2 = st.columns(2)
                    with bcol1:
                        if md_path and md_path.exists():
                            st.download_button(
                                "📝 Download .md",
                                data=md_path.read_text(encoding="utf-8"),
                                file_name=md_path.name,
                                mime="text/markdown",
                                key=f"md_{entry['id']}"
                            )
                    with bcol2:
                        if docx_path and docx_path.exists():
                            st.download_button(
                                "📘 Download .docx",
                                data=docx_path.read_bytes(),
                                file_name=docx_path.name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"docx_{entry['id']}"
                            )
